import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Example script to generate MCQ exams with a custom header
def create_exam(filename, questions, with_correction=False, logo_path=None):
    doc = docx.Document()
    
    # Configure the first page header
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    header = section.first_page_header
    
    # Add logo if provided
    if logo_path and os.path.exists(logo_path):
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = header_para.add_run()
        r.add_picture(logo_path, width=Inches(2.0))
    
    # Add required header text
    header_table = header.add_table(rows=5, cols=2, width=Inches(6.0))
    cells = header_table.cell(0, 0).merge(header_table.cell(0, 1))
    cells.text = "Infohas attached"
    cells.paragraphs[0].runs[0].bold = True
    cells.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    header_table.cell(1, 0).text = "Student Name:"
    header_table.cell(1, 1).text = "________________________"
    header_table.cell(2, 0).text = "Module number:"
    header_table.cell(2, 1).text = "XX (Module Name)"
    header_table.cell(3, 0).text = "Nature of Examination:"
    header_table.cell(3, 1).text = "________________________"
    header_table.cell(4, 0).text = "Date:"
    header_table.cell(4, 1).text = "________________________"
    
    doc.add_paragraph()
    
    title = "MCQ Exam: Module XX"
    if with_correction:
        title += " (Correction Key)"
    doc.add_heading(title, 1)
    
    for i, item in enumerate(questions):
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(item["q"])
        run.bold = True
        
        for j, opt in enumerate(item["options"]):
            prefix = chr(65 + j) + ") "
            opt_p = doc.add_paragraph(prefix + opt)
            if with_correction and j == item["a"]:
                opt_run = opt_p.runs[0]
                opt_run.font.color.rgb = RGBColor(0, 128, 0)
                opt_run.bold = True
                opt_p.add_run(" [CORRECT ANSWER]").bold = True
                opt_p.runs[-1].font.color.rgb = RGBColor(0, 128, 0)
                
    doc.save(filename)
    print(f"Created {filename}")
